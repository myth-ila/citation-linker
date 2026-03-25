# Citation Linker - Converts academic citations in Word documents to hyperlinks
# Author: Mithila Hegde
# AI Assistance: Anthropic's Claude Opus 4.6
# Version: 1.0
# Date: March 2026

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
from typing import List, Tuple

class CitationLinker:
    def __init__(self, input_path: str, output_path: str):
        self.doc = Document(input_path)
        self.output_path = output_path
        self.bib_entries = {}
        self.bib_start_idx = None
        
    def extract_bibliography(self):
        """Extract all bibliography entries"""
        for i, para in enumerate(self.doc.paragraphs):
            if "References" in para.text or "Bibliography" in para.text:
                self.bib_start_idx = i + 1
                break
        
        if self.bib_start_idx is None:
            print("Bibliography section not found")
            return
        
        for para in self.doc.paragraphs[self.bib_start_idx:]:
            if not para.text.strip():
                continue
            
            year_match = re.search(r'\((\d{4})\)', para.text)
            if not year_match:
                continue
            
            year = year_match.group(1)
            text_before_year = para.text[:year_match.start()].strip().rstrip('.,')
            
            last_names = self._parse_bib_authors(text_before_year)
            
            if last_names:
                key = (tuple(last_names), year)
                self.bib_entries[key] = para
                print(f"Found bibliography entry: {last_names} ({year})")
                self._create_bookmark(para, f"{' '.join(last_names)}, {year}")
    
    def _parse_bib_authors(self, text_before_year: str) -> List[str]:
        """
        Parse author last names from bibliography text before the year.
        
        Handles formats like:
          - Silverman, Josh
          - Ries, Al and Jack Trout
          - Chakravarty, Anindita, Alok Kumar, and Rajdeep Grewal
          - Van Alstyne, Marshall W., Geoffrey G. Parker, and Sangeet Paul Choudary
          - Whitler, Kimberly A.
        """
        last_names = []
        
        # Split on " and " to separate the last author
        and_parts = re.split(r'\s+and\s+', text_before_year)
        
        for i, part in enumerate(and_parts):
            part = part.strip().rstrip(',')
            
            if i == 0:
                # First segment: "LastName, FirstName, FirstName2 LastName2,..."
                # or "Van Alstyne, Marshall W., Geoffrey G. Parker"
                # Split on commas
                comma_parts = [p.strip() for p in part.split(',')]
                
                if not comma_parts:
                    continue
                
                # First comma part is always the first author's last name
                first_last = comma_parts[0].strip()
                if first_last:
                    last_names.append(first_last)
                
                # Remaining comma parts: skip the first name of the first author (index 1),
                # then every subsequent part could be a first name of another author
                # Pattern: after "LastName, FirstName", additional authors appear as
                # "FirstName [M.] LastName" in the comma-separated list
                j = 1
                while j < len(comma_parts):
                    segment = comma_parts[j].strip()
                    if not segment:
                        j += 1
                        continue
                    
                    # Check if this segment looks like "FirstName [MiddleInitial] LastName"
                    # i.e., it contains a last name (the last word that starts with uppercase)
                    words = segment.split()
                    
                    if j == 1:
                        # This is the first name of the first author — skip it
                        # But it might also contain additional info if there's no comma
                        # between first name and next author. Usually it's just "FirstName"
                        # or "FirstName M." so we skip.
                        j += 1
                        continue
                    
                    # For j >= 2, this should be "FirstName [M.] LastName"
                    # The last name is the last word that is a full name (not an initial)
                    if words:
                        # Find the last word that isn't an initial (like "G." or "W.")
                        last_name = None
                        for w in reversed(words):
                            if not re.match(r'^[A-Z]\.$', w) and len(w) > 1:
                                last_name = w
                                break
                        if last_name and last_name not in last_names:
                            last_names.append(last_name)
                    
                    j += 1
            else:
                # After "and": "FirstName [M.] LastName" or just "LastName"
                words = part.split()
                if words:
                    # Last non-initial word is the last name
                    last_name = None
                    for w in reversed(words):
                        if not re.match(r'^[A-Z]\.$', w) and len(w) > 1:
                            last_name = w
                            break
                    if last_name and last_name not in last_names:
                        last_names.append(last_name)
        
        return last_names
    
    def _create_bookmark(self, para, bookmark_name):
        """Create a named anchor at the start of the paragraph"""
        try:
            bookmark_name = bookmark_name.replace(" ", "_").replace(",", "").replace(".", "")
            para_elem = para._element
            
            bookmark_start = parse_xml(
                f'<w:bookmarkStart {nsdecls("w")} w:id="0" w:name="{bookmark_name}"/>'
            )
            bookmark_end = parse_xml(
                f'<w:bookmarkEnd {nsdecls("w")} w:id="0"/>'
            )
            
            para_elem.insert(0, bookmark_start)
            para_elem.append(bookmark_end)
        except Exception as e:
            print(f"Could not create bookmark: {e}")
    
    def find_citations(self, para_text: str) -> List[Tuple[str, int, Tuple]]:
        """Find all citations in text, handling semicolon-separated citations."""
        citations = []
        
        bracket_pattern = r'\(([^)]+)\)'
        
        for bracket_match in re.finditer(bracket_pattern, para_text):
            bracket_content = bracket_match.group(1)
            bracket_start = bracket_match.start()
            
            # Split on semicolons to handle multiple citations in one parenthetical
            # e.g., "e.g., Corstjens and Doyle 1989; Whitler 2021"
            # e.g., "DM-LFM; Pang, Liu, and Xu 2022"
            segments = [s.strip() for s in bracket_content.split(';')]
            
            for segment in segments:
                # Find year in this segment
                year_match = re.search(r'(\d{4})', segment)
                if not year_match:
                    continue
                
                year = year_match.group(1)
                before_year = segment[:year_match.start()].strip().rstrip(', ')
                
                # Clean prefixes like "e.g.,", "see", "cf.,", "see also"
                before_year_clean = re.sub(
                    r'^(e\.g\.,?|see|cf\.,?|see also)\s*', '', before_year
                ).strip()
                
                if not before_year_clean:
                    continue
                
                # Build the citation text as it should appear
                citation_text = f"{before_year_clean} {year}"
                
                # Extract author last names from the citation for matching
                # Citation format: "Author1, Author2, and Author3 YEAR"
                # or "Author1 and Author2 YEAR"
                cited_authors = self._parse_citation_authors(before_year_clean)
                
                if not cited_authors:
                    continue
                
                # Match against bibliography entries
                matching_key = self._match_citation_to_bib(cited_authors, year)
                
                if matching_key:
                    # Find the actual position of this citation text in the full paragraph
                    citation_pos = para_text.find(citation_text)
                    if citation_pos == -1:
                        # Try finding within the bracket
                        citation_pos = bracket_start + 1 + bracket_content.find(
                            citation_text.split()[0]
                        )
                    citations.append((citation_text, citation_pos, matching_key))
                    print(f"Found citation: {citation_text}")
        
        return citations
    
    def _parse_citation_authors(self, text: str) -> Tuple:
        """
        Parse author last names from an in-text citation string.
        
        Handles:
          - "Ries and Trout"
          - "Chakravarty, Kumar, and Grewal"
          - "Van Alstyne, Parker, and Choudary"
          - "Silverman"
          - "Whitler"
        """
        authors = []
        
        # Split on " and "
        and_parts = re.split(r'\s+and\s+', text)
        
        for part in and_parts:
            # Each part may contain comma-separated authors
            # e.g., "Chakravarty, Kumar" or "Van Alstyne, Parker"
            comma_parts = [p.strip() for p in part.split(',') if p.strip()]
            
            for cp in comma_parts:
                # Each comma part is a single author name in the citation
                # Could be "Van Alstyne" (multi-word) or "Kumar" (single word)
                # In citations, these are always last names
                if cp:
                    authors.append(cp)
        
        return tuple(authors)
    
    def _match_citation_to_bib(self, cited_authors: Tuple, year: str):
        """
        Match citation authors to a bibliography entry.
        
        Uses flexible matching: each cited author must match (or be contained in)
        a bibliography author, and the year must match.
        """
        for (bib_authors, bib_year), _ in self.bib_entries.items():
            if bib_year != year:
                continue
            
            if len(cited_authors) != len(bib_authors):
                continue
            
            # Check if each cited author matches the corresponding bib author
            all_match = True
            for cited, bib in zip(cited_authors, bib_authors):
                if cited != bib:
                    all_match = False
                    break
            
            if all_match:
                return (bib_authors, bib_year)
        
        return None

    def add_hyperlinks(self):
        """Add hyperlinks using XML manipulation"""
        for para_idx, para in enumerate(self.doc.paragraphs[:self.bib_start_idx - 1]):
            text = para.text
            citations = self.find_citations(text)
            
            if not citations:
                continue
            
            # Sort by position (reverse to process from end to start)
            citations.sort(key=lambda x: x[1], reverse=True)
            
            for citation_text, citation_start, key in citations:
                authors, year = key
                anchor_name = (
                    f"{' '.join(authors)}, {year}".replace(" ", "_").replace(",", "").replace(".", "")
                )
                
                self._replace_citation_with_hyperlink(para, citation_text, anchor_name)
    
    def _replace_citation_with_hyperlink(self, para, citation_text, anchor_name):
        """Replace citation text with hyperlinked version using XML"""
        para_elem = para._element
        
        full_text = para.text
        citation_pos = full_text.find(citation_text)
        
        if citation_pos == -1:
            print(f"Could not find citation '{citation_text}' in paragraph")
            return
        
        # Find which run(s) contain this citation
        char_count = 0
        start_run_idx = None
        
        for run_idx, run in enumerate(para.runs):
            run_start = char_count
            run_end = char_count + len(run.text)
            
            if run_start <= citation_pos < run_end:
                start_run_idx = run_idx
                break
            
            char_count = run_end
        
        if start_run_idx is None:
            print(f"Could not locate citation '{citation_text}' in runs")
            return
        
        # Calculate offset within the starting run
        run = para.runs[start_run_idx]
        offset = citation_pos - sum(len(para.runs[i].text) for i in range(start_run_idx))
        
        # Check if citation spans multiple runs
        citation_end_pos = citation_pos + len(citation_text)
        run_end_pos = sum(len(para.runs[i].text) for i in range(start_run_idx + 1))
        
        if citation_end_pos <= run_end_pos:
            # Citation is within a single run
            before = run.text[:offset]
            after = run.text[offset + len(citation_text):]
            
            hyperlink_xml = (
                f'<w:hyperlink {nsdecls("w")} w:anchor="{anchor_name}">'
                f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
                f'<w:t>{citation_text}</w:t></w:r></w:hyperlink>'
            )
            hyperlink_elem = parse_xml(hyperlink_xml)
            
            run_elem = run._element
            run_parent = run_elem.getparent()
            run_idx_in_parent = list(run_parent).index(run_elem)
            
            run_parent.remove(run_elem)
            
            insert_idx = run_idx_in_parent
            
            if before:
                before_run = parse_xml(
                    f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{before}</w:t></w:r>'
                )
                run_parent.insert(insert_idx, before_run)
                insert_idx += 1
            
            run_parent.insert(insert_idx, hyperlink_elem)
            insert_idx += 1
            
            if after:
                after_run = parse_xml(
                    f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{after}</w:t></w:r>'
                )
                run_parent.insert(insert_idx, after_run)
        else:
            # Citation spans multiple runs — collect and replace them
            runs_to_replace = []
            accumulated = 0
            for i in range(start_run_idx, len(para.runs)):
                runs_to_replace.append(i)
                accumulated = sum(len(para.runs[j].text) for j in range(start_run_idx, i + 1))
                total_from_start = sum(
                    len(para.runs[j].text) for j in range(i + 1)
                )
                if total_from_start >= citation_end_pos:
                    break
            
            # Text before citation in first run
            first_run = para.runs[runs_to_replace[0]]
            before = first_run.text[:offset]
            
            # Text after citation in last run
            last_run = para.runs[runs_to_replace[-1]]
            last_run_start = sum(
                len(para.runs[j].text) for j in range(runs_to_replace[-1])
            )
            after_offset = citation_end_pos - last_run_start
            after = last_run.text[after_offset:]
            
            hyperlink_xml = (
                f'<w:hyperlink {nsdecls("w")} w:anchor="{anchor_name}">'
                f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
                f'<w:t>{citation_text}</w:t></w:r></w:hyperlink>'
            )
            hyperlink_elem = parse_xml(hyperlink_xml)
            
            run_parent = first_run._element.getparent()
            first_elem = first_run._element
            insert_pos = list(run_parent).index(first_elem)
            
            # Remove all affected runs
            for i in runs_to_replace:
                run_parent.remove(para.runs[i]._element)
            
            idx = insert_pos
            if before:
                before_run = parse_xml(
                    f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{before}</w:t></w:r>'
                )
                run_parent.insert(idx, before_run)
                idx += 1
            
            run_parent.insert(idx, hyperlink_elem)
            idx += 1
            
            if after:
                after_run = parse_xml(
                    f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{after}</w:t></w:r>'
                )
                run_parent.insert(idx, after_run)
    
    def save(self):
        """Save the document"""
        self.doc.save(self.output_path)
        print(f"Document saved to {self.output_path}")


if __name__ == '__main__':
    input_file = r"your_document_name.docx"
    output_file = r"your_document_name_linked.docx"
    
    linker = CitationLinker(input_file, output_file)
    linker.extract_bibliography()
    linker.add_hyperlinks()
    linker.save()